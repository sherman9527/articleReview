"""Parse documents into plain text + basic metadata."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ParsedDocument:
    title: str = ""
    text: str = ""
    paragraphs: list[str] = field(default_factory=list)
    word_count: int = 0
    page_count: int = 0
    file_format: str = ""
    file_name: str = ""
    file_size_bytes: int = 0


# ---------------------------------------------------------------------------
# Format parsers
# ---------------------------------------------------------------------------

def _parse_txt(path: Path) -> ParsedDocument:
    encoding = _detect_encoding(path)
    text = path.read_text(encoding=encoding, errors="replace")
    return _build_doc(path, text)


def _parse_md(path: Path) -> ParsedDocument:
    return _parse_txt(path)  # Markdown is plain text


def _parse_docx(path: Path) -> ParsedDocument:
    try:
        import docx
    except ImportError:
        raise RuntimeError(
            "需要安装 python-docx 来解析 .docx 文件。\n"
            "运行: pip install python-docx"
        )

    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    parsed = _build_doc(path, text)
    parsed.paragraphs = paragraphs
    return parsed


def _parse_pdf(path: Path) -> ParsedDocument:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError(
            "需要安装 PyMuPDF 来解析 .pdf 文件。\n"
            "运行: pip install PyMuPDF"
        )

    from . import config

    doc = fitz.open(str(path))
    total_pages = len(doc)
    max_pages = getattr(config, 'MAX_PAGES', 0) or total_pages
    start_page = getattr(config, 'START_PAGE', 1)
    pages_text: list[str] = []
    for i, page in enumerate(doc, 1):
        if i < start_page:
            continue
        if i >= start_page + max_pages:
            break
        page_text = page.get_text()
        if page_text.strip():
            cleaned = _clean_pdf_page(page_text)
            if cleaned.strip():
                pages_text.append(f"【第{i}页】\n{cleaned}")
    doc.close()

    text = "\n\n".join(pages_text)
    parsed = _build_doc(path, text)
    parsed.page_count = total_pages
    return parsed


def _clean_pdf_page(text: str) -> str:
    """Remove garbled lines from PDF-extracted text.

    PyMuPDF sometimes produces garbage for footnotes, formulas, page headers,
    and poorly-encoded fonts. We filter out lines that look like encoding
    artifacts so they don't pollute the analysis.
    """
    cleaned_lines = []
    for line in text.split("\n"):
        line_s = line.strip()
        if not line_s:
            cleaned_lines.append("")
            continue

        # Skip very short lines that are just page numbers / footnote markers
        if len(line_s) <= 2 and not re.search(r'[\u4e00-\u9fff]', line_s):
            continue

        # Count character types
        total = len(line_s)
        chinese = len(re.findall(r'[\u4e00-\u9fff]', line_s))
        latin = len(re.findall(r'[a-zA-Z]', line_s))
        digits = len(re.findall(r'\d', line_s))
        punct = len(re.findall(r'[\u3000-\u303f\uff00-\uffef，。；：？！、""''（）【】《》]', line_s))
        readable = chinese + latin + digits + punct

        # Drop lines where less than 40% of chars are readable
        # (these are typically garbled symbol sequences from bad font encoding)
        if total >= 8 and readable / total < 0.40:
            continue

        # Drop lines that are mostly non-printable or control-range characters
        garbage = len(re.findall(r'[\x00-\x1f\x7f-\x9f]', line_s))
        if garbage > total * 0.2:
            continue

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_PARSERS = {
    ".txt": _parse_txt,
    ".md": _parse_md,
    ".markdown": _parse_md,
    ".docx": _parse_docx,
    ".pdf": _parse_pdf,
}

SUPPORTED_FORMATS = list(_PARSERS.keys())


def parse_document(path: str | Path) -> ParsedDocument:
    """Parse *path* into a :class:`ParsedDocument`."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"文件不存在: {p}")

    suffix = p.suffix.lower()
    parser = _PARSERS.get(suffix)
    if parser is None:
        raise ValueError(
            f"不支持的文件格式: {suffix}\n"
            f"支持的格式: {', '.join(SUPPORTED_FORMATS)}"
        )
    return parser(p)


def _build_doc(path: Path, text: str) -> ParsedDocument:
    text = text.strip()
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]

    # Guess title from first non-empty line (skip page markers like 【第X页】)
    title = ""
    for line in text.split("\n"):
        line = line.strip().lstrip("#").strip()
        if line and not re.match(r'^【第\d+页】$', line):
            title = line[:80]
            break

    # Chinese character count (rough)
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    total_chars = len(text.replace(" ", "").replace("\n", ""))
    word_count = chinese_chars if chinese_chars > total_chars // 2 else total_chars

    return ParsedDocument(
        title=title,
        text=text,
        paragraphs=paragraphs,
        word_count=word_count,
        file_format=path.suffix.lower(),
        file_name=path.name,
        file_size_bytes=path.stat().st_size,
    )


def _detect_encoding(path: Path) -> str:
    """Detect file encoding; fall back to utf-8."""
    try:
        import chardet
        raw = path.read_bytes()[:10000]
        result = chardet.detect(raw)
        if result and result.get("encoding"):
            return result["encoding"]
    except ImportError:
        pass
    return "utf-8"
